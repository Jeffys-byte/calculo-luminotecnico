import streamlit as st
import pandas as pd
import math
import io
import datetime
import base64
import sqlite3
import random

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="Luminotécnica Profissional",
    page_icon="💡",
    layout="wide"
)

# --- CONFIGURAÇÃO DO BANCO DE DADOS SQLITE ---
def init_db():
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            email TEXT PRIMARY KEY,
            senha TEXT NOT NULL,
            criacao TEXT NOT NULL,
            tipo TEXT NOT NULL,
            assinante INTEGER DEFAULT 0,
            token_recuperacao TEXT,
            token_expira TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS clientes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email_usuario TEXT,
            nome TEXT,
            email_cli TEXT,
            telefone TEXT,
            cidade TEXT,
            FOREIGN KEY (email_usuario) REFERENCES usuarios(email)
        )
    ''')
    conn.commit()
    
    cursor.execute("SELECT * FROM usuarios WHERE email = ?", ("jefkar27@gmail.com",))
    if not cursor.fetchone():
        data_criacao = (datetime.datetime.now() - datetime.timedelta(days=30)).strftime("%Y-%m-%d %H:%M:%S")
        cursor.execute("INSERT INTO usuarios (email, senha, criacao, tipo, assinante) VALUES (?, ?, ?, ?, ?)",
                       ("jefkar27@gmail.com", "123", data_criacao, "admin", 1))
        cursor.execute("INSERT INTO clientes (email_usuario, nome, email_cli, telefone, cidade) VALUES (?, ?, ?, ?, ?)",
                       ("jefkar27@gmail.com", "Cliente Geral", "contato@clientegeral.com", "(21) 99999-9999", "Rio de Janeiro - RJ"))
        conn.commit()
    conn.close()

init_db()

def carregar_usuario_db(email):
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    cursor.execute("SELECT senha, criacao, tipo, assinante, token_recuperacao FROM usuarios WHERE email = ?", (email,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return None
    
    senha, criacao_str, tipo, assinante, token_rec = row
    try:
        criacao = datetime.datetime.strptime(criacao_str, "%Y-%m-%d %H:%M:%S")
    except:
        criacao = datetime.datetime.now()
        
    cursor.execute("SELECT nome, email_cli, telefone, cidade FROM clientes WHERE email_usuario = ?", (email,))
    clientes_rows = cursor.fetchall()
    banco_clientes = [{"Nome": r[0], "Email": r[1], "Telefone": r[2], "Cidade": r[3]} for r in clientes_rows]
    if not banco_clientes:
        banco_clientes = [{"Nome": "Cliente Geral", "Email": "contato@clientegeral.com", "Telefone": "(21) 99999-9999", "Cidade": "Rio de Janeiro - RJ"}]

    conn.close()
    return {
        "senha": senha,
        "criacao": criacao,
        "tipo": tipo,
        "assinante": bool(assinante),
        "token_recuperacao": token_rec,
        "banco_clientes": banco_clientes
    }

def salvar_usuario_db(email, senha, tipo="cliente", assinante=0):
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    data_criacao = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute("INSERT OR REPLACE INTO usuarios (email, senha, criacao, tipo, assinante) VALUES (?, ?, ?, ?, ?)",
                   (email, senha, data_criacao, tipo, assinante))
    cursor.execute("INSERT INTO clientes (email_usuario, nome, email_cli, telefone, cidade) VALUES (?, ?, ?, ?, ?)",
                   (email, "Cliente Exemplo", "exemplo@email.com", "(21) 98888-8888", "Rio de Janeiro - RJ"))
    conn.commit()
    conn.close()

def atualizar_senha_db(email, nova_senha):
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    cursor.execute("UPDATE usuarios SET senha = ?, token_recuperacao = NULL WHERE email = ?", (nova_senha, email))
    conn.commit()
    conn.close()

def excluir_usuario_db(email):
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    # Remove primeiro os clientes vinculados ao usuário para manter a integridade
    cursor.execute("DELETE FROM clientes WHERE email_usuario = ?", (email,))
    # Remove o usuário
    cursor.execute("DELETE FROM usuarios WHERE email = ?", (email,))
    conn.commit()
    conn.close()

def gerar_token_recuperacao(email):
    token = str(random.randint(100000, 999999))
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    cursor.execute("UPDATE usuarios SET token_recuperacao = ? WHERE email = ?", (token, email))
    conn.commit()
    conn.close()
    return token

def adicionar_cliente_db(email_usuario, nome, email_cli, telefone, cidade):
    conn = sqlite3.connect('luminotecnica.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO clientes (email_usuario, nome, email_cli, telefone, cidade) VALUES (?, ?, ?, ?, ?)",
                   (email_usuario, nome, email_cli, telefone, cidade))
    conn.commit()
    conn.close()

def listar_todos_usuarios_db():
    conn = sqlite3.connect('luminotecnica.db')
    df_users = pd.read_sql_query("SELECT email, criacao, tipo, assinante FROM usuarios", conn)
    df_cli = pd.read_sql_query("SELECT email_usuario, nome, email_cli, telefone, cidade FROM clientes", conn)
    conn.close()
    return df_users, df_cli

def definir_fundo_personalizado_base64(img_bytes=None, url_imagem=None):
    if img_bytes:
        encoded = base64.b64encode(img_bytes).decode()
        bg_val = f"url('data:image/jpeg;base64,{encoded}')"
    elif url_imagem:
        bg_val = f"url('{url_imagem}')"
    else:
        return
        
    css = f"""
    <style>
    .stApp {{
        background-image: linear-gradient(rgba(10, 15, 30, 0.75), rgba(10, 15, 30, 0.85)), {bg_val};
        background-size: cover;
        background-position: center;
        background-repeat: no-repeat;
        background-attachment: fixed;
    }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def verificar_autenticacao():
    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False
    if "usuario_email" not in st.session_state:
        st.session_state.usuario_email = None

    if not st.session_state.autenticado:
        definir_fundo_personalizado_base64(url_imagem="https://images.unsplash.com/photo-1600585154340-be6161a56a0c?q=80&w=1920&auto=format&fit=crop")
        
        st.markdown("<h2 style='text-align: center; color: #ffffff;'>🔐 Área Restrita - Luminotécnica Profissional</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #cbd5e0;'>Crie sua conta e ganhe <b>24 horas de teste gratuito</b>, ou faça login se já tiver cadastro.</p>", unsafe_allow_html=True)
        
        tab_login, tab_cadastro, tab_planos = st.tabs(["🔑 Fazer Login", "📝 Criar Conta Grátis (Teste 24h)", "💳 Assinar (R$ 19,90/mês)"])
        
        with tab_login:
            with st.form("form_login"):
                email_input = st.text_input("E-mail cadastrado", value="").strip().lower()
                senha_input = st.text_input("Senha", type="password", value="").strip()
                btn_entrar = st.form_submit_button("Entrar no Sistema")
                
                if btn_entrar:
                    user_data = carregar_usuario_db(email_input)
                    if user_data and user_data["senha"] == senha_input:
                        agora = datetime.datetime.now()
                        tempo_criacao = user_data["criacao"]
                        horas_decorridas = (agora - tempo_criacao).total_seconds() / 3600
                        
                        if user_data["tipo"] == "admin" or horas_decorridas <= 24 or user_data.get("assinante", False):
                            st.session_state.autenticado = True
                            st.session_state.usuario_email = email_input
                            st.success("Login realizado com sucesso!")
                            st.rerun()
                        else:
                            st.error("⏰ Seu período de teste de 24 horas expirou. Vá na aba 'Assinar' para continuar usando por apenas R$ 19,90/mês.")
                    else:
                        st.error("E-mail ou senha incorretos, ou usuário não encontrado.")

            with st.expander("🔑 Esqueci / Redefinir minha senha com Token"):
                st.markdown("Insira seu e-mail para solicitar um **Código de Segurança (Token)** de redefinição.")
                with st.form("form_pedir_token"):
                    email_rec = st.text_input("E-mail cadastrado para recuperação").strip().lower()
                    btn_gerar_token = st.form_submit_button("Gerar Código de Segurança")
                    
                    if btn_gerar_token:
                        if email_rec:
                            if carregar_usuario_db(email_rec):
                                token_criado = gerar_token_recuperacao(email_rec)
                                st.success(f"Código de segurança gerado com sucesso! (Simulação de envio): **{token_criado}**")
                            else:
                                st.error("Este e-mail não está cadastrado.")
                        else:
                            st.error("Informe o e-mail.")

                st.markdown("---")
                st.markdown("Já possui o código? Digite-o abaixo junto com a nova senha:")
                with st.form("form_confirmar_token"):
                    email_conf = st.text_input("Confirme seu e-mail").strip().lower()
                    token_dig = st.text_input("Código de Segurança (Token de 6 dígitos)").strip()
                    nova_senha_rec = st.text_input("Digite a nova senha", type="password").strip()
                    btn_redefinir = st.form_submit_button("Validar e Atualizar Senha")
                    
                    if btn_redefinir:
                        if email_conf and token_dig and nova_senha_rec:
                            usr_check = carregar_usuario_db(email_conf)
                            if usr_check:
                                if usr_check["token_recuperacao"] == token_dig:
                                    atualizar_senha_db(email_conf, nova_senha_rec)
                                    st.success("Senha redefinida com segurança! Faça login na aba ao lado com a nova senha.")
                                else:
                                    st.error("Código de segurança (Token) inválido ou incorreto.")
                            else:
                                st.error("E-mail não encontrado.")
                        else:
                            st.error("Preencha todos os campos.")

        with tab_cadastro:
            st.markdown("### ⚡ Comece a usar agora mesmo")
            with st.form("form_cadastro"):
                novo_email = st.text_input("Seu E-mail principal", value="").strip().lower()
                nova_senha = st.text_input("Crie uma Senha", type="password", value="").strip()
                btn_cadastrar = st.form_submit_button("Criar Conta e Iniciar Teste Grátis")
                
                if btn_cadastrar:
                    if novo_email and nova_senha:
                        if carregar_usuario_db(novo_email):
                            st.warning("Este e-mail já está cadastrado. Faça login na primeira aba.")
                        else:
                            salvar_usuario_db(novo_email, nova_senha, tipo="cliente", assinante=0)
                            st.session_state.autenticado = True
                            st.session_state.usuario_email = novo_email
                            st.success("Conta criada com sucesso! Seu teste de 24 horas começou.")
                            st.rerun()
                    else:
                        st.error("Preencha todos os campos para criar a conta.")

        with tab_planos:
            st.markdown("### 🚀 Assinatura Profissional")
            st.markdown("Tenha acesso ilimitado a todos os cálculos normativos (NBR ISO/CIE 8995-1).")
            
            st.markdown("""
                <div style="background-color: #1a202c; border: 2px solid #ffc107; padding: 25px; border-radius: 8px; text-align: center; color: white; margin-top: 15px; margin-bottom: 15px;">
                    <h3 style="margin-top: 0; color: #ffc107;">Desbloqueie o Acesso Completo</h3>
                    <p style="font-size: 16px; color: #cbd5e0; margin-bottom: 5px;">Tenha relatórios ilimitados e sem restrições por apenas:</p>
                    <div style="font-size: 34px; font-weight: bold; color: #28a745; margin: 15px 0;">R$ 19,90 / mês</div>
                </div>
            """, unsafe_allow_html=True)
            
            st.link_button("💳 Pagar com Mercado Pago (R$ 19,90)", "https://mpago.la/2sbQvQ9", use_container_width=True)
                    
        return False

    return True

if not verificar_autenticacao():
    st.stop()

# --- VERIFICAÇÃO CONTÍNUA DO TEMPO DE TESTE ---
email_atual = st.session_state.usuario_email
user_info_atual = carregar_usuario_db(email_atual)

if user_info_atual and user_info_atual["tipo"] != "admin" and not user_info_atual.get("assinante", False):
    agora = datetime.datetime.now()
    horas_decorridas = (agora - user_info_atual["criacao"]).total_seconds() / 3600
    if horas_decorridas > 24:
        st.session_state.autenticado = False
        st.error("⏰ Seu período de teste de 24 horas expirou. Faça a assinatura para continuar utilizando o sistema.")
        st.stop()

banco_clientes_usuario = user_info_atual["banco_clientes"] if user_info_atual else [{"Nome": "Cliente Geral", "Email": "contato@clientegeral.com", "Telefone": "(21) 99999-9999", "Cidade": "Rio de Janeiro - RJ"}]

TABELA_NORMA = {
    "Residências - Salas de Estar / Dormitórios": 150,
    "Residências - Cozinhas / Banheiros": 300,
    "Escritórios - Geral / Digitação": 500,
    "Escritórios - Reunião / Conferência": 300,
    "Comércio - Lojas de Departamento / Varejo": 500,
    "Comércio - Supermercados / Áreas de Circulação": 300,
    "Indústria - Montagem Grosso (Ex: Mecânica Pesada)": 200,
    "Indústria - Montagem Média (Ex: Eletrônicos)": 500,
    "Escolas - Salas de Aula / Laboratórios": 300,
    "Hospitais - Enfermarias": 100,
    "Garagens - Áreas de Estacionamento / Circulação": 75,
}

if "banco_luminarias" not in st.session_state:
    st.session_state.banco_luminarias = [
        {"Fabricante": "Ecolume", "Modelo": "Painel LED 24W Redondo Sobrepor", "Lumens": 1920, "Potencia": 24.0, "Tipo": "Painel/Luminária"},
        {"Fabricante": "Philips", "Modelo": "Painel LED 18W Quadrado Embutir", "Lumens": 1440, "Potencia": 18.0, "Tipo": "Painel/Luminária"},
        {"Fabricante": "Philips", "Modelo": "Painel LED 24W Redondo Embutir", "Lumens": 1920, "Potencia": 24.0, "Tipo": "Painel/Luminária"},
        {"Fabricante": "Osram", "Modelo": "Luminária LED Estanque 36W", "Lumens": 3600, "Potencia": 36.0, "Tipo": "Painel/Luminária"},
        {"Fabricante": "Taschibra", "Modelo": "Painel LED Slim 12W Quadrado", "Lumens": 960, "Potencia": 12.0, "Tipo": "Painel/Luminária"},
    ]

if "banco_fitas" not in st.session_state:
    st.session_state.banco_fitas = [
        {"Fabricante": "Gaya", "Modelo": "Fita LED 10W/m IP20", "Lumens": 900, "Potencia": 10.0},
        {"Fabricante": "Super LED", "Modelo": "Fita LED 14.4W/m SMD5050", "Lumens": 1200, "Potencia": 14.4},
    ]

def gerar_docx_consolidado(dados_cliente, dados_profissional, lista_ambientes, logo_file=None):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(9.5)
    style_normal.font.color.rgb = RGBColor(50, 50, 50)

    HEX_COR_PRIMARIA = "1A365D"
    COR_TEXTO_TITULO = RGBColor(26, 54, 93)

    def set_cell_background(cell, hex_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    if logo_file:
        try:
            doc.add_picture(logo_file, width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass

    data_atual_str = datetime.date.today().strftime("%d/%m/%Y")
    p_t = doc.add_paragraph()
    r_t = p_t.add_run("RELATÓRIO LUMINOTÉCNICO EXECUTIVO CONSOLIDADO")
    r_t.bold = True
    r_t.font.size = Pt(14)
    r_t.font.color.rgb = COR_TEXTO_TITULO

    p_sub = doc.add_paragraph()
    p_sub.add_run(f"Cliente / Empreendimento: {dados_cliente.get('Nome', 'Cliente Geral')} | Método dos Lúmens\n")
    p_sub.add_run(f"Responsável Técnico: {dados_profissional.get('nome', 'Não informado')} — Registro: {dados_profissional.get('registro', 'Não informado')} | Data de Emissão: {data_atual_str}\n")
    p_sub.add_run(f"Norma de Referência: NBR ISO/CIE 8995-1 & NBR 5410")
    for r in p_sub.runs:
        r.font.size = Pt(9)

    doc.add_paragraph()

    for idx, amb in enumerate(lista_ambientes):
        if idx > 0:
            doc.add_page_break()

        p_amb = doc.add_paragraph()
        r_amb = p_amb.add_run(f"AMBIENTE: {amb['nome'].upper()}")
        r_amb.bold = True
        r_amb.font.size = Pt(11.5)
        r_amb.font.color.rgb = COR_TEXTO_TITULO

        t1 = doc.add_table(rows=15, cols=4)
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        t1.autofit = False
        w1 = [Inches(2.5), Inches(1.0), Inches(1.8), Inches(1.2)]

        headers1 = ["Parâmetro", "Símbolo", "Valor Adotado", "Unidade"]
        for ci, h in enumerate(headers1):
            cell = t1.cell(0, ci)
            cell.width = w1[ci]
            set_cell_background(cell, HEX_COR_PRIMARIA)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)

        dados_bloco1 = [
            ("Nome do Ambiente", "—", amb['nome'], "—"),
            ("Comprimento / Largura / Pé-Direito", "C x L x H", f"{amb['comp']:.2f} x {amb['larg']:.2f} x {amb['pe_direito']:.2f}", "m"),
            ("Área Total do Piso", "A", f"{amb['area']:.2f}", "m²"),
            ("Índice do Recinto (Geometria)", "k", f"{amb['k_indice']:.2f}", "—"),
            ("Iluminância Requerida (Normativa)", "Ereq", f"{amb['lux_req']:.2f}", "lx"),
            ("Fatores de Utilização e Depreciação", "u / d", f"u = {amb['fator_u']:.2f} | d = {amb['fator_d']:.2f}", "—"),
            ("Fonte Luminosa / Equipamento", "Φ", f"{amb['fluxo_unidade_rel']:,.2f} lm".replace(",", "."), amb['modelo_lum']),
            ("Quantidade de Equipamentos Adotada", "N", f"{amb['qtd_real_str']}", amb['unidade_medida_qtd']),
            ("Arranjo Luminoso Distribuído", "—", f"{amb['arranjo_str']}", "arr."),
            ("Espaçamentos e Afastamentos (C / L)", "dc / dl", f"Entre: {amb['dist_c_entre']:.2f}m / {amb['dist_l_entre']:.2f}m | Paredes: {amb['dist_c_parede']:.2f}m / {amb['dist_l_parede']:.2f}m", "m"),
            ("Iluminância Real Alcançada", "Ereal", f"{amb['lux_real']:.2f}", "lx"),
            ("Potência Total Instalada e DPI", "P / DPI", f"{amb['pot_total']:.2f} W | {amb['dpi']:.2f} W/m²", "W / W/m²"),
            ("Status Final de Conformidade", "—", "CONFORME (Aprovado)" if amb['conforme'] else "NÃO CONFORME", "—"),
            ("Anotações / Responsabilidade", "ART/CREA", f"Prof: {dados_profissional.get('nome', 'N/A')} ({dados_profissional.get('registro', 'N/A')})", "—")
        ]

        for ri, row_vals in enumerate(dados_bloco1):
            row_cells = t1.rows[ri + 1].cells
            bg = "F7FAFC" if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(row_vals):
                cell = row_cells[ci]
                cell.width = w1[ci]
                set_cell_background(cell, bg)
                set_cell_margins(cell)
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(8.5)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def gerar_pdf_consolidado(dados_cliente, dados_profissional, lista_ambientes, logo_file=None):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=24, leftMargin=24, topMargin=30, bottomMargin=30)
    story = []
    styles = getSampleStyleSheet()

    cor_primaria = colors.HexColor("#1A365D")
    
    titulo_style = ParagraphStyle(
        'TituloRelatorio',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=13,
        textColor=cor_primaria,
        spaceAfter=4
    )
    
    texto_style = ParagraphStyle(
        'TextoNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        textColor=colors.HexColor("#323232")
    )

    th_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        textColor=colors.whitesmoke,
        alignment=0
    )

    td_left = ParagraphStyle(
        'TableCellLeft',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        textColor=colors.HexColor("#323232"),
        alignment=0
    )

    story.append(Paragraph("RELATÓRIO LUMINOTÉCNICO EXECUTIVO CONSOLIDADO", titulo_style))
    data_atual_str = datetime.date.today().strftime("%d/%m/%Y")
    
    info_txt = f"<b>Cliente / Empreendimento:</b> {dados_cliente.get('Nome', 'Cliente Geral')} | Método dos Lúmens<br/>" \
               f"<b>Responsável Técnico:</b> {dados_profissional.get('nome', 'Não informado')} — Registro: {dados_profissional.get('registro', 'Não informado')} | Data: {data_atual_str}<br/>" \
               f"<b>Norma de Referência:</b> NBR ISO/CIE 8995-1 & NBR 5410"
    story.append(Paragraph(info_txt, texto_style))
    story.append(Spacer(1, 10))

    for idx, amb in enumerate(lista_ambientes):
        if idx > 0:
            story.append(PageBreak())

        story.append(Paragraph(f"<b>AMBIENTE: {amb['nome'].upper()}</b>", titulo_style))
        story.append(Spacer(1, 6))

        tabela_bruta = [
            ["Parâmetro", "Símbolo", "Valor Adotado", "Unidade"],
            ["Nome do Ambiente", "—", amb['nome'], "—"],
            ["Comprimento / Largura / Pé-Direito", "C x L x H", f"{amb['comp']:.2f} x {amb['larg']:.2f} x {amb['pe_direito']:.2f}", "m"],
            ["Área Total do Piso", "A", f"{amb['area']:.2f}", "m²"],
            ["Índice do Recinto (Geometria)", "k", f"{amb['k_indice']:.2f}", "—"],
            ["Iluminância Requerida (Normativa)", "Ereq", f"{amb['lux_req']:.2f} lx", "NBR ISO/CIE 8995-1"],
            ["Fatores de Utilização e Depreciação", "u / d", f"u = {amb['fator_u']:.2f} | d = {amb['fator_d']:.2f}", "—"],
            ["Fonte Luminosa / Equipamento", "Φ", f"{amb['fluxo_unidade_rel']:,.2f} lm", amb['modelo_lum']],
            ["Quantidade de Equipamentos Adotada", "N", f"{amb['qtd_real_str']}", amb['unidade_medida_qtd']],
            ["Arranjo Luminoso Distribuído", "—", f"{amb['arranjo_str']}", "arr."],
            ["Espaçamentos (Entre e Paredes)", "dc / dl", f"Entre: {amb['dist_c_entre']:.2f}m / {amb['dist_l_entre']:.2f}m | Paredes: {amb['dist_c_parede']:.2f}m / {amb['dist_l_parede']:.2f}m", "m"],
            ["Iluminância Real Alcançada", "Ereal", f"{amb['lux_real']:.2f} lx", "Calculado"],
            ["Potência Total Instalada e DPI", "P / DPI", f"{amb['pot_total']:.2f} W | {amb['dpi']:.2f} W/m²", "W / W/m²"],
            ["Status Final de Conformidade", "—", "CONFORME (Aprovado)" if amb['conforme'] else "NÃO CONFORME", "—"],
            ["Responsabilidade Técnica", "ART", f"{dados_profissional.get('nome', 'N/A')} ({dados_profissional.get('registro', 'N/A')})", "—"]
        ]

        tabela_dados = []
        for r_idx, row in enumerate(tabela_bruta):
            nova_linha = []
            for c_idx, cell in enumerate(row):
                if r_idx == 0:
                    nova_linha.append(Paragraph(str(cell), th_style))
                else:
                    nova_linha.append(Paragraph(str(cell), td_left))
            tabela_dados.append(nova_linha)

        t = Table(tabela_dados, colWidths=[200, 60, 180, 124])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), cor_primaria),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor("#F7FAFC"), colors.white])
        ]))
        
        story.append(t)
        story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

st.title("💡 Luminotécnica Profissional")

# --- BARRA DE STATUS E BOTÃO DE UPGRADE DIRETO PARA CONTAS DE TESTE ---
if user_info_atual and user_info_atual["tipo"] != "admin" and not user_info_atual.get("assinante", False):
    agora = datetime.datetime.now()
    horas_decorridas = (agora - user_info_atual["criacao"]).total_seconds() / 3600
    horas_restantes = max(0, 24 - horas_decorridas)
    
    col_aviso, col_botao = st.columns([3, 1])
    with col_aviso:
        st.info(f"⏳ **Conta em Período de Teste Gratuito** | Restam aprox. **{horas_restantes:.1f} horas** de acesso.")
    with col_botao:
        st.link_button("🚀 Virar PRO (R$ 19,90)", "https://mpago.la/2sbQvQ9", use_container_width=True)

st.markdown(f"**Sessão Ativa:** {st.session_state.get('usuario_email', 'Usuário')}")

if st.sidebar.button("🚪 Sair do Sistema"):
    st.session_state.autenticado = False
    st.rerun()

if user_info_atual and user_info_atual["tipo"] == "admin":
    st.sidebar.markdown("---")
    with st.sidebar.expander("👑 Painel Admin ℹ️"):
        st.markdown("Gerencie os registros do banco de dados SQLite de usuários e clientes cadastrados.")
        df_u, df_c = listar_todos_usuarios_db()
        st.markdown("**Usuários:**")
        st.dataframe(df_u, use_container_width=True)
        st.markdown("**Clientes:**")
        st.dataframe(df_c, use_container_width=True)
        
        st.markdown("---")
        st.markdown("🎁 **Gerenciar Usuário / Acessos / Exclusão**")
        email_alvo = st.selectbox("Selecione o e-mail do usuário", df_u["email"].tolist() if not df_u.empty else [], key="admin_email_alvo")
        
        col_pa1, col_pa2 = st.columns(2)
        with col_pa1:
            if st.button("🔄 Renovar Teste (24h)"):
                conn = sqlite3.connect('luminotecnica.db')
                cursor = conn.cursor()
                agora_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                cursor.execute("UPDATE usuarios SET criacao = ? WHERE email = ?", (agora_str, email_alvo))
                conn.commit()
                conn.close()
                st.success(f"Teste de {email_alvo} renovado por mais 24h!")
                st.rerun()
                
        with col_pa2:
            if st.button("⭐ Liberar como PRO"):
                conn = sqlite3.connect('luminotecnica.db')
                cursor = conn.cursor()
                cursor.execute("UPDATE usuarios SET assinante = 1 WHERE email = ?", (email_alvo,))
                conn.commit()
                conn.close()
                st.success(f"Usuário {email_alvo} promovido a Assinante PRO com acesso livre!")
                st.rerun()

        if st.button("🚫 Revogar PRO (Voltar para Padrão/Expirado)", use_container_width=True):
            conn = sqlite3.connect('luminotecnica.db')
            cursor = conn.cursor()
            cursor.execute("UPDATE usuarios SET assinante = 0 WHERE email = ?", (email_alvo,))
            conn.commit()
            conn.close()
            st.warning(f"Status PRO de {email_alvo} revogado com sucesso!")
            st.rerun()

        st.markdown("---")
        # Botão com confirmação para exclusão de usuário
        if st.button("🗑️ Excluir Usuário Selecionado do Banco", use_container_width=True):
            if email_alvo == "jefkar27@gmail.com":
                st.error("Não é permitido excluir o usuário Administrador principal do sistema!")
            else:
                excluir_usuario_db(email_alvo)
                st.success(f"Usuário {email_alvo} excluído com sucesso do banco de dados!")
                st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("### 🏢 Logotipo do Projeto ℹ️")
logo_upload = st.sidebar.file_uploader("Enviar Logo (.png, .jpg)", type=["png", "jpg", "jpeg"], help="Envie o logotipo da sua empresa para customizar os relatórios finais.")

st.sidebar.markdown("---")
st.sidebar.markdown("### 👷 Identificação do Profissional ℹ️")
prof_nome = st.sidebar.text_input("Nome do Profissional", value="", key="prof_nome_input", help="Nome que aparecerá como responsável técnico no relatório.")
prof_registro = st.sidebar.text_input("Registro / CREA / CAU", value="", key="prof_reg_input", help="Número do registro profissional (CREA/CAU).")
prof_celular = st.sidebar.text_input("Celular / WhatsApp", value="", key="prof_cel_input", help="Contato para o rodapé do projeto.")
prof_email = st.sidebar.text_input("E-mail Profissional", value="", key="prof_email_input", help="E-mail corporativo do projetista.")

st.markdown("---")

aba_cli_mod, aba_calc_mod, aba_fitas_mod = st.tabs([
    "📇 Cadastro e Seleção de Clientes", 
    "🏠 1. Cálculo de Luminárias e Painéis", 
    "✨ 2. Projeto de Fitas LED (PRO)"
])

with aba_cli_mod:
    st.markdown("### 📇 Cadastro e Seleção de Clientes ℹ️")
    with st.expander("➕ Cadastrar Novo Cliente no Sistema"):
        with st.form("form_novo_cliente"):
            col_nc1, col_nc2 = st.columns(2)
            with col_nc1:
                cad_nome_cli = st.text_input("Nome / Razão Social do Cliente")
                cad_tel_cli = st.text_input("Telefone / WhatsApp")
            with col_nc2:
                cad_email_cli = st.text_input("E-mail do Cliente")
                cad_cidade_cli = st.text_input("Cidade / Estado (Ex: Rio de Janeiro - RJ)")
                
            btn_salvar_cliente = st.form_submit_button("Salvar Cliente")
            if btn_salvar_cliente:
                if cad_nome_cli.strip() != "":
                    adicionar_cliente_db(
                        email_atual, 
                        cad_nome_cli, 
                        cad_email_cli if cad_email_cli else "Não informado",
                        cad_tel_cli if cad_tel_cli else "Não informado",
                        cad_cidade_cli if cad_cidade_cli else "Não informado"
                    )
                    st.success(f"Cliente '{cad_nome_cli}' cadastrado com sucesso!")
                    st.rerun()
                else:
                    st.error("O nome do cliente é obrigatório.")

    lista_nomes_clientes = [c["Nome"] for c in banco_clientes_usuario]
    cliente_selecionado_nome = st.selectbox("Selecione o Cliente para este Projeto ℹ️", lista_nomes_clientes, help="Escolha qual cliente/obra receberá os cálculos atuais.")
    cliente_dados_obj = next((c for c in banco_clientes_usuario if c["Nome"] == cliente_selecionado_nome), banco_clientes_usuario[0])

st.markdown("---")

with aba_calc_mod:
    st.markdown(f"### 🛋️ Ambientes para o Cliente: **{banco_clientes_usuario[0]['Nome'] if 'banco_clientes_usuario' in locals() else 'Cliente'}**")

    with st.expander("⚙️ Cadastrar Nova Luminária no Banco ℹ️"):
        st.markdown("Adicione novas luminárias comerciais com seus fluxos em lúmens e potências para uso rápido.")
        with st.form("form_nova_lum"):
            col_fl1, col_fl2, col_fl3, col_fl4, col_fl5 = st.columns(5)
            with col_fl1:
                novo_tipo = st.selectbox("Categoria", ["Painel/Luminária", "Industrial"])
            with col_fl2:
                novo_fab = st.text_input("Fabricante", value="Philips")
            with col_fl3:
                novo_mod = st.text_input("Modelo", value="Painel LED")
            with col_fl4:
                novo_lum = st.number_input("Fluxo (lm)", value=1440.0, step=50.0)
            with col_fl5:
                nova_pot = st.number_input("Potência (W)", value=18.0, step=1.0)
                
            btn_salvar_lum = st.form_submit_button("Salvar no Banco")
            if btn_salvar_lum:
                st.session_state.banco_luminarias.append({
                    "Fabricante": novo_fab,
                    "Modelo": novo_mod,
                    "Lumens": novo_lum,
                    "Potencia": nova_pot,
                    "Tipo": novo_tipo
                })
                st.success("Luminária cadastrada com sucesso!")

    if "ambientes" not in st.session_state:
        st.session_state.ambientes = [{"id": 1, "nome": "Ambiente 1"}]

    col_add, col_rem = st.columns([1, 1])
    with col_add:
        if st.button("➕ Adicionar Novo Ambiente ℹ️", help="Adiciona um novo espaço/cômodo independente para cálculo luminotécnico."):
            novo_id = st.session_state.ambientes[-1]["id"] + 1 if st.session_state.ambientes else 1
            st.session_state.ambientes.append({"id": novo_id, "nome": f"Ambiente {novo_id}"})
            st.rerun()
    with col_rem:
        if len(st.session_state.ambientes) > 1 and st.button("🗑️ Remover Último Ambiente"):
            st.session_state.ambientes.pop()
            st.rerun()

    lista_calculos_ambientes = []

    for amb_atual in st.session_state.ambientes:
        with st.container():
            st.markdown(f"#### 📐 Ambiente: {amb_atual['nome']}")
            
            col_n1, col_n2 = st.columns(2)
            with col_n1:
                novo_nome = st.text_input("Nome do Ambiente", value=amb_atual['nome'], key=f"nome_amb_{amb_atual['id']}", help="Identificação descritiva do cômodo (ex: Sala de Estar).")
            with col_n2:
                tipo_atividade = st.selectbox("Atividade / Norma (NBR ISO/CIE 8995-1) ℹ️", list(TABELA_NORMA.keys()), key=f"ativ_{amb_atual['id']}", help="Define o nível de iluminância mínima em Lux exigido pela norma técnica para cada tipo de uso.")

            lux_padrao_norma = TABELA_NORMA[tipo_atividade]
            col_lux1, col_lux2 = st.columns(2)
            with col_lux1:
                usar_lux_manual = st.checkbox("Alterar Iluminância (Lux) Manualmente? ℹ️", key=f"chk_lux_{amb_atual['id']}", help="Permite definir um valor personalizado de Lux em vez do padrão normativo.")
            with col_lux2:
                if usar_lux_manual:
                    lux_req = st.number_input("Iluminância Desejada (lx)", value=float(lux_padrao_norma), step=10.0, key=f"lux_man_{amb_atual['id']}")
                else:
                    lux_req = float(lux_padrao_norma)
                    st.markdown(f"**Iluminância Normativa:** {lux_req} lx")

            st.markdown("##### 💡 Seleção de Luminária / Painel")
            banco_ativo = st.session_state.banco_luminarias
            opcoes_banco_str = [f"{l['Fabricante']} - {l['Modelo']} ({l['Lumens']} lm / {l['Potencia']} W)" for l in banco_ativo]
            opcoes_banco_str.append("⚙️ Inserir Manual / Personalizado")
            
            escolha_banco = st.selectbox("Selecionar Equipamento ℹ️", opcoes_banco_str, key=f"lum_escolha_{amb_atual['id']}", help="Escolha uma luminária do banco cadastrado ou insira características customizadas.")

            if escolha_banco != "⚙️ Inserir Manual / Personalizado":
                idx_escolhido = opcoes_banco_str.index(escolha_banco)
                lum_sel = banco_ativo[idx_escolhido]
                fluxo_base, potencia_base = lum_sel["Lumens"], lum_sel["Potencia"]
                modelo_desc_relatorio = f"[Painel/Luminária] {lum_sel['Fabricante']} - {lum_sel['Modelo']}"
            else:
                col_m1, col_m2 = st.columns(2)
                with col_m1:
                    fluxo_base = st.number_input("Fluxo Luminoso da Unidade (lm)", value=1920.0, step=50.0, key=f"fluxo_man_lum_{amb_atual['id']}")
                with col_m2:
                    potencia_base = st.number_input("Potência Unitária (W)", value=24.0, step=1.0, key=f"pot_man_lum_{amb_atual['id']}")
                modelo_desc_relatorio = "[Painel/Luminária] Personalizado"

            st.markdown("##### Geometria e Fatores ℹ️")
            col_g1, col_g2, col_g3 = st.columns(3)
            with col_g1:
                comp = st.number_input("Comprimento (m)", value=3.0, step=0.1, key=f"comp_{amb_atual['id']}", help="Comprimento em metros do recinto.")
                pe_direito = st.number_input("Pé-Direito (m)", value=2.9, step=0.1, key=f"pd_{amb_atual['id']}", help="Altura total do piso ao teto.")
            with col_g2:
                larg = st.number_input("Largura (m)", value=2.0, step=0.1, key=f"larg_{amb_atual['id']}", help="Largura em metros do recinto.")
                hp = st.number_input("Plano de Trabalho (m)", value=0.75, step=0.05, key=f"hp_{amb_atual['id']}", help="Altura do plano onde ocorre a tarefa visual (padrão de escritório: 0,75m).")
            with col_g3:
                hp_desc = st.number_input("Rebaixamento / Suspensão (m)", value=0.0, step=0.05, key=f"hdesc_{amb_atual['id']}", help="Distância de rebaixamento do teto (ex: forro de gesso estruturado ou luminária pendente).")

            col_f1, col_f2 = st.columns(2)
            with col_f1:
                fator_u = st.slider("Fator de Utilização (u) ℹ️", 0.3, 0.8, 0.70, 0.05, key=f"fu_{amb_atual['id']}", help="Eficiência luminosa da sala baseada nas refletâncias das paredes, teto e piso.")
            with col_f2:
                fator_d = st.slider("Fator de Depreciação (d) ℹ️", 0.5, 0.9, 0.80, 0.05, key=f"fd_{amb_atual['id']}", help="Fator de manutenção que considera o acúmulo de poeira e envelhecimento das lâmpadas.")

            area = comp * larg
            hu = pe_direito - hp - hp_desc
            k_indice = (comp * larg) / (hu * (comp + larg)) if hu > 0 else 1.0
            fluxo_req = (lux_req * area) / (fator_u * fator_d) if (fator_u * fator_d) > 0 else 0

            fluxo_lampada = fluxo_base
            potencia_lampada = potencia_base
            qtd_teorica = fluxo_req / fluxo_lampada if fluxo_lampada > 0 else 0
            qtd_min_sugerida = math.ceil(qtd_teorica)
            if qtd_min_sugerida < 1:
                qtd_min_sugerida = 1

            st.markdown("##### 🎛️ Configuração do Arranjo e Quantidade")
            usar_qtd_manual = st.checkbox("Definir quantidade total de luminárias manualmente? ℹ️", key=f"chk_qtd_man_{amb_atual['id']}", help="Permite ignorar o cálculo automático de quantidade e fixar o número exato de luminárias desejado.")

            if usar_qtd_manual:
                qtd_real = st.number_input("Quantidade Total de Luminárias", min_value=1, value=qtd_min_sugerida, step=1, key=f"qtd_manual_val_{amb_atual['id']}")
                linhas_man, colunas_man = 1, int(qtd_real)
                arranjo_str = f"Arranjo Livre ({int(qtd_real)} unidades)"
            else:
                col_arr1, col_arr2 = st.columns(2)
                with col_arr1:
                    linhas_man = st.number_input("Quantidade de Linhas", min_value=1, value=1, step=1, key=f"linhas_man_{amb_atual['id']}")
                with col_arr2:
                    colunas_man = st.number_input("Quantidade de Colunas", min_value=1, value=2, step=1, key=f"colunas_man_{amb_atual['id']}")
                
                qtd_real = linhas_man * colunas_man
                arranjo_str = f"{linhas_man} Linhas x {colunas_man} Colunas"

            if qtd_real < qtd_min_sugerida:
                st.warning(f"⚠️ O quantitativo selecionado ({qtd_real} un) está abaixo do mínimo teórico calculado ({qtd_min_sugerida} un).")

            relacao_sh_padrao = 1.25 
            espacamento_max_permitido = hu * relacao_sh_padrao

            calc_c_auto = comp / colunas_man if colunas_man > 0 else comp
            calc_l_auto = larg / linhas_man if linhas_man > 0 else larg

            usar_espacamento_manual = st.checkbox("Definir afastamentos e espaçamentos manualmente? ℹ️", key=f"chk_esp_{amb_atual['id']}", help="Ajusta de forma personalizada a distância entre as luminárias e as paredes.")

            if usar_espacamento_manual:
                col_em1, col_em2 = st.columns(2)
                with col_em1:
                    dist_c_entre = st.number_input("Distância entre Luminárias (Comprimento - m)", value=float(calc_c_auto), step=0.1, key=f"dist_c_man_{amb_atual['id']}")
                    dist_c_parede = st.number_input("Afastamento da Parede (Comprimento - m)", value=float(calc_c_auto/2.0), step=0.1, key=f"dc_par_man_{amb_atual['id']}")
                with col_em2:
                    dist_l_entre = st.number_input("Distância entre Luminárias (Largura - m)", value=float(calc_l_auto), step=0.1, key=f"dist_l_man_{amb_atual['id']}")
                    dist_l_parede = st.number_input("Afastamento da Parede (Largura - m)", value=float(calc_l_auto/2.0), step=0.1, key=f"dl_par_man_{amb_atual['id']}")
            else:
                dist_c_entre = calc_c_auto
                dist_c_parede = dist_c_entre / 2.0
                dist_l_entre = calc_l_auto
                dist_l_parede = dist_l_entre / 2.0

            espacamento_critico_excedido = (dist_c_entre > espacamento_max_permitido) or (dist_l_entre > espacamento_max_permitido)

            fluxo_instalado = qtd_real * fluxo_lampada
            pot_total = qtd_real * potencia_lampada
            
            lux_real = (fluxo_instalado * fator_u * fator_d) / area if area > 0 else 0
            dpi = pot_total / area if area > 0 else 0
            variacao_fluxo_pct = ((fluxo_instalado - fluxo_req) / fluxo_req) * 100 if fluxo_req > 0 else 0
            conforme = lux_real >= lux_req

            st.success(f"✨ **Conferência de Quantidade:** **{int(qtd_real)} unidades** adotadas no arranjo ({area:.1f} m²).")
            st.info(f"📏 **Espaçamentos Práticos:** Entre: C={dist_c_entre:.2f}m / L={dist_l_entre:.2f}m | **Máximo Recomendado (S/H):** {espacamento_max_permitido:.2f}m")
            
            if espacamento_critico_excedido:
                st.warning(f"⚠️ **Atenção:** O espaçamento adotado entre as luminárias excede o limite técnico recomendado para esta altura de montagem ($Hu$ = {hu:.2f}m).")

            lista_calculos_ambientes.append({
                "id": amb_atual["id"],
                "nome": novo_nome,
                "comp": comp,
                "larg": larg,
                "pe_direito": pe_direito,
                "hp": hp,
                "hp_desc": hp_desc,
                "hu": hu,
                "area": area,
                "k_indice": k_indice,
                "lux_req": lux_req,
                "fluxo_req": fluxo_req,
                "fluxo_unidade_rel": fluxo_lampada,
                "pot_unidade_rel": potencia_lampada,
                "unidade_pot_desc": "Consumo Unitário (W)",
                "fator_u": fator_u,
                "fator_d": fator_d,
                "desc_utilizacao": "Ambiente residencial / Padrão",
                "desc_depreciacao": "Limpeza periódica / Padrão",
                "modelo_lum": modelo_desc_relatorio,
                "fluxo_instalado": fluxo_instalado,
                "qtd_teorica": qtd_teorica,
                "qtd_real_str": str(int(qtd_real)),
                "unidade_medida_qtd": "un",
                "arranjo_str": arranjo_str,
                "dist_c_entre": dist_c_entre,
                "dist_c_parede": dist_c_parede,
                "dist_l_entre": dist_l_entre,
                "dist_l_parede": dist_l_parede,
                "lux_real": lux_real,
                "pot_total": pot_total,
                "dpi": dpi,
                "variacao_fluxo_pct": variacao_fluxo_pct,
                "conforme": conforme
            })
            st.markdown("---")

with aba_fitas_mod:
    st.markdown("### ✨ Projeto de Fitas LED Lineares ℹ️")
    col_fita_1, col_fita_2 = st.columns(2)
    with col_fita_1:
        fita_comp = st.number_input("Comprimento Linear da Sanca / Perfil (m)", value=10.0, step=0.5, key="fita_comp_m", help="Metragem linear total onde a fita LED será instalada.")
        fita_lux_req = st.number_input("Iluminância Alvo (lx)", value=200.0, step=10.0, key="fita_lux_alvo", help="Nível de iluminação desejado para o projeto linear.")
    with col_fita_2:
        banco_fita_opcoes = [f"{f['Fabricante']} - {f['Modelo']} ({f['Lumens']} lm/m)" for f in st.session_state.banco_fitas]
        banco_fita_opcoes.append("⚙️ Personalizada")
        sel_fita_aba = st.selectbox("Escolher Fita LED", banco_fita_opcoes, key="sel_fita_aba_key")
        
        if "Personalizada" not in sel_fita_aba:
            idx_f = banco_fita_opcoes.index(sel_fita_aba)
            lm_metro = st.session_state.banco_fitas[idx_f]["Lumens"]
        else:
            lm_metro = st.number_input("Fluxo por Metro (lm/m)", value=900.0, step=50.0, key="fita_lm_man")

    metragem_calculada_fita = (fita_lux_req * fita_comp) / lm_metro if lm_metro > 0 else 0
    st.info(f"📏 **Metragem Teórica Calculada de Fita LED:** **{metragem_calculada_fita:.2f} metros lineares**.")

st.subheader("3. Emissão de Relatório Luminotécnico ℹ️")

is_admin_or_subscriber = (user_info_atual and (user_info_atual.get("tipo") == "admin" or user_info_atual.get("assinante", False)))

if not is_admin_or_subscriber:
    st.warning("🔒 **Recurso Exclusivo para Assinantes:** O período de teste permite realizar os cálculos na tela, mas a geração e o download dos relatórios oficiais (.docx e .pdf) exigem uma assinatura ativa. Vá acima na barra de teste ou faça o upgrade para liberar!")
else:
    col_dl1, col_dl2 = st.columns(2)
    
    with col_dl1:
        if st.button("📄 Gerar Relatório Word (.docx)", use_container_width=True):
            try:
                dados_prof_dict = {
                    "nome": prof_nome if prof_nome else "Não informado",
                    "registro": prof_registro if prof_registro else "Não informado",
                    "celular": prof_celular if prof_celular else "Não informado",
                    "email": prof_email if prof_email else "Não informado"
                }
                logo_bytes = io.BytesIO(logo_upload.getvalue()) if logo_upload is not None else None
                cliente_ativo_rel = banco_clientes_usuario[0] if banco_clientes_usuario else {"Nome": "Cliente Geral"}
                arquivo_docx_bytes = gerar_docx_consolidado(cliente_ativo_rel, dados_prof_dict, lista_ambientes, logo_file=logo_bytes)
                
                st.success("Word gerado com sucesso!")
                st.download_button(
                    label="📥 Baixar Arquivo .docx",
                    data=arquivo_docx_bytes,
                    file_name=f"Relatorio_{cliente_ativo_rel['Nome'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Erro ao gerar Word: {e}")

    with col_dl2:
        if st.button("📑 Gerar Relatório PDF (.pdf)", use_container_width=True):
            try:
                dados_prof_dict = {
                    "nome": prof_nome if prof_nome else "Não informado",
                    "registro": prof_registro if prof_registro else "Não informado",
                    "celular": prof_celular if prof_celular else "Não informado",
                    "email": prof_email if prof_email else "Não informado"
                }
                logo_bytes = io.BytesIO(logo_upload.getvalue()) if logo_upload is not None else None
                cliente_ativo_rel = banco_clientes_usuario[0] if banco_clientes_usuario else {"Nome": "Cliente Geral"}
                arquivo_pdf_bytes = gerar_pdf_consolidado(cliente_ativo_rel, dados_prof_dict, lista_ambientes, logo_file=logo_bytes)
                
                st.success("PDF gerado com sucesso!")
                st.download_button(
                    label="📥 Baixar Arquivo .pdf",
                    data=arquivo_pdf_bytes,
                    file_name=f"Relatorio_{cliente_ativo_rel['Nome'].replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Erro ao gerar PDF: {e}")
